"""
Web server for document upload and processing
Handles file uploads and runs the extraction pipeline
Supports React frontend via CORS
"""

from flask import Flask, render_template, request, jsonify, g
from flask_cors import CORS
from werkzeug.utils import secure_filename
import os
import sys
import shutil
import time
from dotenv import load_dotenv
from integration.integrated_pipeline import IntegratedPipeline
from knowledge_graph.nebula_handler import NebulaGraphClient
from vector_database.weaviate_handler import WeaviateClient
from utils.helpers import print_section
import uuid

load_dotenv()

app = Flask(__name__)
CORS(app, resources={r"/api/*": {"origins": ["http://localhost:5173", "http://localhost:5174", "http://localhost:8080", "http://localhost:5000", "http://127.0.0.1:5173", "http://127.0.0.1:5174"]}})
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['UPLOAD_FOLDER'] = 'uploads'

# Create uploads folder if it doesn't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)


# ── Simple request/response logger using print + flush ─────────
def _log(msg):
    """Print a timestamped log line, flushing immediately."""
    ts = time.strftime('%H:%M:%S')
    print(f"[{ts}] {msg}", flush=True)


@app.before_request
def _before():
    g.req_start = time.time()
    parts = [f">> {request.method} {request.path}"]
    qs = request.query_string.decode('utf-8', errors='replace')
    if qs:
        parts.append(f"?{qs}")
    ct = request.content_type or ''
    if 'json' in ct:
        body = request.get_data(as_text=True)
        preview = (body[:200] + '...') if len(body) > 200 else body
        parts.append(f"  body={preview}")
    elif 'multipart' in ct:
        fnames = [f.filename for f in request.files.values()] if request.files else []
        parts.append(f"  files={fnames}")
    _log(''.join(parts))


@app.after_request
def _after(response):
    dur = (time.time() - g.get('req_start', time.time())) * 1000
    code = response.status_code
    tag = 'OK' if code < 400 else 'ERR'
    line = f"<< [{tag}] {code} {request.method} {request.path} ({dur:.0f}ms)"
    ct = response.content_type or ''
    if 'json' in ct:
        body = response.get_data(as_text=True)
        preview = (body[:300] + '...') if len(body) > 300 else body
        line += f"\n     response={preview}"
    _log(line)
    return response

# Initialize pipeline and handlers
pipeline = IntegratedPipeline(mode="semantic")
nebula_handler = NebulaGraphClient()
weaviate_handler = WeaviateClient()

ALLOWED_EXTENSIONS = {
    'txt', 'pdf', 'docx', 'doc',
    'png', 'jpg', 'jpeg', 'tif', 'tiff', 'bmp'
}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def _env_truthy(value: str | None) -> bool:
    if value is None:
        return False
    return value.strip().lower() in {"1", "true", "yes", "y", "on"}

@app.route('/')
def index():
    """Serve the main page"""
    return render_template('index.html')

@app.route('/api/upload', methods=['POST'])
def upload_file():
    """Handle file upload and process through pipeline"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file provided'}), 400
        
        file = request.files['file']
        
        if not file.filename:
            return jsonify({'success': False, 'error': 'No file selected'}), 400
        
        if not allowed_file(str(file.filename)):
            return jsonify({'success': False, 'error': 'File type not allowed. Use: txt, pdf, docx, doc, png, jpg, jpeg, tif, tiff, bmp'}), 400
        
        # Save file
        filename = secure_filename(str(file.filename))
        unique_filename = f"{uuid.uuid4()}_{filename}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        file.save(filepath)
        
        # Extract text from file
        text_content, extraction_error = extract_text(filepath)

        if extraction_error:
            return jsonify({'success': False, 'error': extraction_error}), 400

        if not text_content or not text_content.strip():
            return jsonify({
                'success': False,
                'error': 'Could not extract text from file. If this is a scanned PDF, upload an image (PNG/JPG) or enable OCR for PDFs.'
            }), 400
        
        # Run pipeline
        print_section(f"Processing: {filename}")
        results = pipeline.run_complete_pipeline(text_content)

        # Get entities from workflow and transform to frontend format (use 'name' instead of 'value')
        raw_entities = results.get('workflow', {}).get('entities', [])
        extracted_entities = [
            {
                'name': e.get('value', e.get('name', '')),  # Backend uses 'value', frontend expects 'name'
                'type': e.get('type', 'CUSTOM'),
                'confidence': e.get('confidence', 0.85)
            }
            for e in raw_entities if e.get('value') or e.get('name')
        ]
        # Get relationships and transform to frontend format
        raw_relationships = results.get('relationships', [])
        def _strip_type_prefix(id_str):
            """Dynamically strip type prefix (e.g. 'person_john_doe' -> 'john doe')"""
            if not id_str:
                return ''
            if '_' in id_str:
                return '_'.join(id_str.split('_')[1:]).replace('_', ' ')
            return id_str
        
        extracted_relationships = [
            {
                'source': r.get('source') or _strip_type_prefix(r.get('from_id', '')),
                'target': r.get('target') or _strip_type_prefix(r.get('to_id', '')),
                'relationship_type': r.get('type', r.get('relationship_type', 'RELATED_TO')),
                'confidence': r.get('confidence', 0.8)
            }
            for r in raw_relationships if r.get('from_id') or r.get('source')
        ]
        
        # Debug the entities and relationships being sent
        print(f"DEBUG: Sending {len(extracted_entities)} entities to frontend")
        print(f"DEBUG: Sending {len(extracted_relationships)} relationships to frontend")
        if extracted_entities:
            print(f"DEBUG: First entity sample: {extracted_entities[0]}")
        if extracted_relationships:
            print(f"DEBUG: First relationship sample: {extracted_relationships[0]}")
        
        # Get vector document UUID (first ID if multiple stored)
        vector_ids = results.get('vector_storage', {}).get('document_ids', [])
        vector_uuid = vector_ids[0] if vector_ids else 'N/A'
        
        # Get graph storage results
        graph_storage = results.get('graph_storage', {})
        entities_added = graph_storage.get('entities_added', 0)
        relationships_added = graph_storage.get('relationships_added', 0)
        
        # Debug output
        print(f"DEBUG: entities_added={entities_added}, relationships_added={relationships_added}")
        print(f"DEBUG: extracted_relationships={len(extracted_relationships)}")
        
        return jsonify({
            'success': True,
            'filename': filename,
            'text_preview': text_content[:500] + '...' if len(text_content) > 500 else text_content,
            'results': {
                'entities_count': entities_added,
                'relationships_count': relationships_added,
                'vector_document_uuid': vector_uuid,
                'workflow_status': results.get('workflow', {}).get('status', 'unknown'),
                'entities': extracted_entities,
                'relationships': extracted_relationships,
                'errors': results.get('workflow', {}).get('errors', [])
            }
        }), 200
    
    except Exception as e:
        print(f"Error: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

def _tesseract_available() -> bool:
    """Return True if the Tesseract OCR engine appears available on PATH."""
    return shutil.which('tesseract') is not None


def _detect_tesseract_cmd() -> str | None:
    """Return an executable path for tesseract if found.

    Priority:
      1) TESSERACT_CMD env var
      2) PATH (shutil.which)
      3) Common Windows install locations
    """
    env_cmd = os.getenv('TESSERACT_CMD')
    if env_cmd and os.path.exists(env_cmd):
        return env_cmd

    on_path = shutil.which('tesseract')
    if on_path:
        return on_path

    candidates = [
        os.path.join(os.environ.get('ProgramFiles', ''), 'Tesseract-OCR', 'tesseract.exe'),
        os.path.join(os.environ.get('ProgramFiles(x86)', ''), 'Tesseract-OCR', 'tesseract.exe'),
        os.path.join(os.environ.get('LocalAppData', ''), 'Programs', 'Tesseract-OCR', 'tesseract.exe'),
        os.path.join(os.environ.get('LocalAppData', ''), 'Tesseract-OCR', 'tesseract.exe'),
    ]
    for candidate in candidates:
        if candidate and os.path.exists(candidate):
            return candidate

    return None


def _pdf_ocr_enabled() -> bool:
    """Whether OCR fallback for PDFs is enabled."""
    return _env_truthy(os.getenv('ENABLE_PDF_OCR'))


def _pdf_ocr_max_pages() -> int:
    """Max number of pages to OCR from a PDF (avoid long-running requests)."""
    raw = os.getenv('PDF_OCR_MAX_PAGES', '10')
    try:
        value = int(raw)
        return max(1, min(value, 200))
    except ValueError:
        return 10


def _pdf_ocr_dpi() -> int:
    """Rasterization DPI used before OCR."""
    raw = os.getenv('PDF_OCR_DPI', '200')
    try:
        value = int(raw)
        return max(72, min(value, 600))
    except ValueError:
        return 200


def _ocr_pdf(filepath: str) -> tuple[str | None, str | None]:
    """OCR a PDF by rasterizing pages and running Tesseract.

    Returns:
        (text, error_message)
    """
    try:
        try:
            import fitz  # type: ignore[import-not-found]  # PyMuPDF
        except ImportError:
            return None, "PDF OCR requires PyMuPDF. Install with: pip install pymupdf"

        try:
            from PIL import Image
            import pytesseract
        except ImportError:
            return None, "PDF OCR requires pillow + pytesseract. Install with: pip install pillow pytesseract"

        tesseract_cmd = _detect_tesseract_cmd()
        if not tesseract_cmd:
            return None, (
                "PDF OCR requires the Tesseract engine installed. Install it and either restart the terminal "
                "so PATH refreshes, or set TESSERACT_CMD to the full path of tesseract.exe."
            )

        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

        lang = os.getenv('PDF_OCR_LANG', 'eng')
        max_pages = _pdf_ocr_max_pages()
        dpi = _pdf_ocr_dpi()

        text_chunks: list[str] = []

        with fitz.open(filepath) as doc:
            page_count = min(len(doc), max_pages)
            if page_count <= 0:
                return "", None

            zoom = dpi / 72.0
            matrix = fitz.Matrix(zoom, zoom)

            for i in range(page_count):
                page = doc.load_page(i)
                pix = page.get_pixmap(matrix=matrix, alpha=False)
                image = Image.frombytes('RGB', (pix.width, pix.height), pix.samples)
                page_text = pytesseract.image_to_string(image, lang=lang)
                if page_text and page_text.strip():
                    text_chunks.append(page_text)

        return "\n\n".join(text_chunks), None

    except Exception as e:
        return None, f"PDF OCR failed: {str(e)}"


def extract_text(filepath):
    """Extract text from various file formats.

    Returns:
        (text, error_message)
    """
    try:
        lower_path = filepath.lower()

        if lower_path.endswith('.txt'):
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read(), None
        
        elif lower_path.endswith('.pdf'):
            try:
                import PyPDF2
                with open(filepath, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ''.join((page.extract_text() or '') for page in reader.pages)
                    if text and text.strip():
                        return text, None

                    if _pdf_ocr_enabled():
                        ocr_text, ocr_error = _ocr_pdf(filepath)
                        if ocr_error:
                            return None, ocr_error
                        return ocr_text or "", None

                    return text, None
            except ImportError:
                return None, "PDF support requires PyPDF2. Install with: pip install PyPDF2"
        
        elif lower_path.endswith(('.docx', '.doc')):
            try:
                from docx import Document
                doc = Document(filepath)
                text = '\n'.join(para.text for para in doc.paragraphs)
                return text, None
            except ImportError:
                return None, "DOCX support requires python-docx. Install with: pip install python-docx"

        elif lower_path.endswith(('.png', '.jpg', '.jpeg', '.tif', '.tiff', '.bmp')):
            try:
                from PIL import Image
                import pytesseract

                tesseract_cmd = _detect_tesseract_cmd()
                if not tesseract_cmd:
                    return None, (
                        "OCR requires the Tesseract engine installed. Install it and either restart the terminal "
                        "so PATH refreshes, or set TESSERACT_CMD to the full path of tesseract.exe."
                    )

                pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

                with Image.open(filepath) as image:
                    text = pytesseract.image_to_string(image)
                    return text, None
            except ImportError:
                return None, "Image OCR requires pillow + pytesseract. Install with: pip install pillow pytesseract"
            except Exception as e:
                return None, f"Image OCR failed: {str(e)}"
        
        return None, "Unsupported file type"
    
    except Exception as e:
        print(f"Text extraction error: {str(e)}")
        return None, f"Text extraction error: {str(e)}"

@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint with service statuses"""
    import requests
    
    weaviate_connected = False
    nebula_connected = False
    
    # Check Weaviate
    try:
        response = requests.get("http://localhost:8080/v1/.well-known/ready", timeout=2)
        weaviate_connected = response.status_code == 200
    except:
        pass
    
    # Check NebulaGraph
    try:
        from nebula3.gclient.net import ConnectionPool
        from nebula3.Config import Config
        config = Config()
        config.max_connection_pool_size = 1
        pool = ConnectionPool()
        if pool.init([('127.0.0.1', 9669)], config):
            session = pool.get_session('root', 'nebula')
            session.release()
            pool.close()
            nebula_connected = True
    except:
        pass
    
    return jsonify({
        'status': 'ok',
        'message': 'Server running',
        'weaviate_connected': weaviate_connected,
        'nebula_connected': nebula_connected
    }), 200


@app.route('/api/process', methods=['POST'])
def process_text():
    """Process text directly without file upload"""
    try:
        data = request.get_json()
        if not data or 'text' not in data:
            return jsonify({'success': False, 'error': 'No text provided'}), 400
        
        text_content = data['text'].strip()
        if not text_content:
            return jsonify({'success': False, 'error': 'Empty text provided'}), 400
        
        # Run pipeline
        print_section("Processing text input")
        results = pipeline.run_complete_pipeline(text_content)

        # Get entities from workflow and transform to frontend format (use 'name' instead of 'value')
        raw_entities = results.get('workflow', {}).get('entities', [])
        extracted_entities = [
            {
                'name': e.get('value', e.get('name', '')),  # Backend uses 'value', frontend expects 'name'
                'type': e.get('type', 'CUSTOM'),
                'confidence': e.get('confidence', 0.85)
            }
            for e in raw_entities if e.get('value') or e.get('name')
        ]
        
        # Get relationships and transform to frontend format
        raw_relationships = results.get('relationships', [])
        def _strip_type_prefix(id_str):
            """Dynamically strip type prefix (e.g. 'person_john_doe' -> 'john doe')"""
            if not id_str:
                return ''
            if '_' in id_str:
                return '_'.join(id_str.split('_')[1:]).replace('_', ' ')
            return id_str
        
        extracted_relationships = [
            {
                'source': r.get('source') or _strip_type_prefix(r.get('from_id', '')),
                'target': r.get('target') or _strip_type_prefix(r.get('to_id', '')),
                'relationship_type': r.get('type', r.get('relationship_type', 'RELATED_TO')),
                'confidence': r.get('confidence', 0.8)
            }
            for r in raw_relationships if r.get('from_id') or r.get('source')
        ]
        
        # Debug the entities and relationships being sent
        print(f"DEBUG: Sending {len(extracted_entities)} entities to frontend")
        print(f"DEBUG: Sending {len(extracted_relationships)} relationships to frontend")
        
        # Get vector document UUID
        vector_ids = results.get('vector_storage', {}).get('document_ids', [])
        vector_uuid = vector_ids[0] if vector_ids else 'N/A'
        
        # Get graph storage results
        graph_storage = results.get('graph_storage', {})
        entities_added = graph_storage.get('entities_added', 0)
        relationships_added = graph_storage.get('relationships_added', 0)
        
        return jsonify({
            'success': True,
            'text_preview': text_content[:500] + '...' if len(text_content) > 500 else text_content,
            'results': {
                'entities_count': entities_added,
                'relationships_count': relationships_added,
                'vector_document_uuid': vector_uuid,
                'workflow_status': results.get('workflow', {}).get('status', 'unknown'),
                'entities': extracted_entities,
                'relationships': extracted_relationships,
                'errors': results.get('workflow', {}).get('errors', [])
            }
        }), 200
    
    except Exception as e:
        print(f"Error processing text: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/search', methods=['GET'])
def semantic_search():
    """Semantic search using Weaviate vector database"""
    try:
        query = request.args.get('q', '')
        limit = int(request.args.get('limit', 10))
        
        if not query:
            return jsonify({'success': False, 'error': 'No query provided'}), 400
        
        # Use the standalone weaviate handler to search
        try:
            results = weaviate_handler.semantic_search(query, limit=limit)
            return jsonify({
                'success': True,
                'results': results
            }), 200
        except Exception as e:
            print(f"Weaviate search error: {str(e)}")
            return jsonify({'success': True, 'results': []}), 200
    
    except Exception as e:
        print(f"Search error: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/semantic-search', methods=['POST'])
def advanced_semantic_search():
    """
    Advanced semantic search using meaning-based matching.
    
    Example queries:
    - "Why are shipments getting delayed?" → Finds port congestion, weather events
    - "Weather related supply chain issues" → Finds heavy rainfall, monsoon impacts
    - "Who works at Company X?" → Returns employees with relationship mapping
    
    Body: {
        "query": "Why are shipments getting delayed?",
        "limit": 10,
        "category": "supply_chain" (optional),
        "include_graph": true (optional)
    }
    
    Returns:
    - results: Semantic search matches with entities and relationships
    - relationship_answer: Explanation of how relationships were mapped
    - relationship_table: Tabular format for chart display
    """
    try:
        data = request.get_json() or {}
        query = data.get('query', '')
        limit = int(data.get('limit', 10))
        category = data.get('category')
        include_graph = data.get('include_graph', True)
        
        if not query:
            return jsonify({'success': False, 'error': 'No query provided'}), 400
        
        try:
            from semantic_search.semantic_engine import SemanticSearchEngine
            from semantic_search.relationship_mapper import EnhancedRelationshipMapper
            
            engine = SemanticSearchEngine()
            relationship_mapper = EnhancedRelationshipMapper()
            
            # Perform semantic search
            results = engine.semantic_search(query, limit=limit, category_filter=category)
            
            # Get query expansion terms for frontend display
            query_expansion = engine._expand_query_semantically(query) if hasattr(engine, '_expand_query_semantically') else []
            
            # Collect all entities and text for relationship mapping
            all_entities = []
            all_text = ""
            for r in results:
                all_entities.extend(r.entities)
                all_text += r.content + "\n"
            
            # Use enhanced relationship mapper
            relationship_answer = relationship_mapper.generate_relationship_answer(
                query=query,
                entities=all_entities,
                text=all_text,
                semantic_matches=[{
                    'content': r.content,
                    'score': r.score,
                    'entities': r.entities
                } for r in results]
            )
            
            response = {
                'success': True,
                'query': query,
                'query_expansion': list(set(query_expansion))[:15],
                'results': [],
                'total_found': len(results),
                # LLM-generated answer from search results
                'llm_answer': '',
                # Relationship-based answer components
                'relationship_answer': {
                    'direct_answer': relationship_answer.direct_answer,
                    'explanation': relationship_answer.relationship_explanation,
                    'graph_summary': relationship_answer.graph_summary
                },
                'relationship_table': relationship_answer.relationship_table,
                'all_relationships': [
                    {
                        'id': f"rel-{i}",
                        'source': r.source,
                        'source_type': r.source_type.upper(),
                        'target': r.target,
                        'target_type': r.target_type.upper(),
                        'type': r.relationship_type,
                        'confidence': r.confidence,
                        'context': r.context_snippet[:100] if r.context_snippet else '',
                        'mapping_reason': r.mapping_reason
                    }
                    for i, r in enumerate(relationship_answer.relationships)
                ]
            }
            
            # Build result items with relationships
            for r in results:
                result_item = {
                    'id': r.id,
                    'document_id': r.metadata.get('document_id', r.id) if r.metadata else r.id,
                    'content': r.content[:500] + '...' if len(r.content) > 500 else r.content,
                    'score': r.score,
                    'highlights': r.highlights,
                    'source_type': r.source_type,
                    'metadata': r.metadata,
                    'entities': [
                        {
                            'id': f"entity-{i}",
                            'name': e.get('value', e.get('name', '')),
                            'type': e.get('type', 'CUSTOM').upper(),
                            'confidence': e.get('confidence', 0.85)
                        }
                        for i, e in enumerate(r.entities) if e
                    ],
                    'relationships': []
                }
                
                # Map relationships for this document's entities using enhanced mapper
                doc_entities = r.entities
                doc_text = r.content
                
                if doc_entities:
                    doc_relationships = relationship_mapper.map_relationships_for_entities(
                        doc_entities, doc_text, query
                    )
                    
                    result_item['relationships'] = [
                        {
                            'id': f"rel-{j}",
                            'source': rel.source,
                            'source_type': rel.source_type.upper(),
                            'target': rel.target,
                            'target_type': rel.target_type.upper(),
                            'type': rel.relationship_type,
                            'confidence': rel.confidence,
                            'context': rel.context_snippet[:80] if rel.context_snippet else '',
                            'mapping_reason': rel.mapping_reason
                        }
                        for j, rel in enumerate(doc_relationships[:15])  # Limit to 15 per doc
                    ]
                
                response['results'].append(result_item)
            
            # Add graph context if requested
            if include_graph and results:
                try:
                    from semantic_search.graph_traversal import GraphTraversal
                    graph = GraphTraversal()
                    
                    graph_context = []
                    unique_entities = {e.get('value', e.get('name', '')): e for e in all_entities if e}
                    
                    for entity_name, entity in list(unique_entities.items())[:5]:
                        entity_type = entity.get('type', 'unknown')
                        try:
                            relationships = graph.get_entity_relationships(entity_name, entity_type)
                            if relationships:
                                graph_context.append({
                                    "entity": entity_name,
                                    "type": entity_type,
                                    "relationships": relationships[:5]
                                })
                        except Exception:
                            pass
                    
                    response['graph_context'] = graph_context
                except Exception as graph_err:
                    print(f"Graph context error: {graph_err}")
                    response['graph_context'] = []

            # Generate LLM-based answer from search results
            if results:
                try:
                    llm_answer = engine.generate_answer_from_results(query, results)
                    response['llm_answer'] = llm_answer
                except Exception as ans_err:
                    print(f"LLM answer generation error: {ans_err}")

            return jsonify(response), 200
            
        except Exception as e:
            print(f"Semantic search error: {str(e)}")
            import traceback
            traceback.print_exc()
            return jsonify({'success': False, 'error': str(e)}), 500
    
    except Exception as e:
        print(f"Search error: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/affected-companies', methods=['POST'])
def find_affected_companies():
    """
    Find companies indirectly affected by an entity.
    
    Example: "Which companies are indirectly affected by Mumbai port congestion?"
    
    Body: {
        "entity": "Mumbai Port",
        "max_hops": 3 (optional)
    }
    """
    try:
        data = request.get_json() or {}
        entity = data.get('entity', data.get('event', ''))  # Support both 'entity' and 'event'
        max_hops = int(data.get('max_hops', 3))
        
        if not entity:
            return jsonify({'success': False, 'error': 'No entity specified'}), 400
        
        try:
            from semantic_search.semantic_pipeline import SemanticGraphPipeline
            pipeline = SemanticGraphPipeline()
            
            result = pipeline.find_affected_companies(entity)
            
            # Format for frontend
            affected_companies = []
            for company in result.get('affected_companies', []):
                affected_companies.append({
                    'name': company.get('name', company.get('company', '')),
                    'connection_type': company.get('connection_type', company.get('relationship', 'RELATED')),
                    'path': company.get('path', []),
                    'impact_score': company.get('impact_score', 0.75)
                })
            
            # Also get semantic context
            from semantic_search.semantic_engine import SemanticSearchEngine
            engine = SemanticSearchEngine()
            semantic_results = engine.semantic_search(entity, limit=3)
            
            semantic_context = [
                {
                    'id': r.id,
                    'document_id': r.id,
                    'content': r.content[:200] + '...' if len(r.content) > 200 else r.content,
                    'score': r.score
                }
                for r in semantic_results
            ]
            
            return jsonify({
                'success': True,
                'entity': entity,
                'affected_companies': affected_companies,
                'semantic_context': semantic_context,
                'total_found': len(affected_companies)
            }), 200
            
        except Exception as e:
            print(f"Affected companies search error: {str(e)}")
            return jsonify({'success': False, 'entity': entity, 'affected_companies': [], 'error': str(e)}), 500
    
    except Exception as e:
        print(f"Error: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/similar-patterns', methods=['POST'])
def find_similar_patterns():
    """
    Find similar incidents/patterns in the knowledge base.
    
    Example: "Find incidents similar to cyber attacks on cloud vendors"
    
    Body: {
        "query": "cyber attacks on cloud vendors",
        "limit": 10
    }
    """
    try:
        data = request.get_json() or {}
        query = data.get('query', data.get('pattern', ''))  # Support both 'query' and 'pattern'
        limit = int(data.get('limit', 10))
        
        if not query:
            return jsonify({'success': False, 'error': 'No query specified'}), 400
        
        try:
            from semantic_search.semantic_pipeline import SemanticGraphPipeline
            pipeline = SemanticGraphPipeline()
            
            result = pipeline.find_similar_patterns(query)
            
            # Format patterns for frontend with safe defaults
            patterns = []
            for pattern in result.get('similar_patterns', []):
                # Safely get values with proper defaults to avoid None comparisons
                content = pattern.get('content') or pattern.get('name') or ''
                category = pattern.get('category') or 'general'
                score = pattern.get('score')
                # Ensure score is a valid number
                try:
                    score = float(score) if score is not None else 0.5
                except (TypeError, ValueError):
                    score = 0.5
                
                # Safely extract entities
                entities_list = pattern.get('entities') or []
                formatted_entities = []
                for i, e in enumerate(entities_list):
                    if isinstance(e, dict):
                        name = e.get('value') or e.get('name') or ''
                        etype = e.get('type') or 'CUSTOM'
                        if name:  # Only add if we have a name
                            formatted_entities.append({
                                'id': f"entity-{i}",
                                'name': str(name),
                                'type': str(etype)
                            })
                
                patterns.append({
                    'content': str(content),
                    'category': str(category),
                    'score': score,
                    'entities': formatted_entities
                })
            
            # Safely sort by score (all scores are now guaranteed to be numbers)
            patterns.sort(key=lambda x: x.get('score', 0), reverse=True)
            
            return jsonify({
                'success': True,
                'query': query,
                'patterns': patterns[:limit],
                'categories': result.get('categories') or {},
                'total_found': len(patterns)
            }), 200
            
        except Exception as e:
            import traceback
            print(f"Pattern search error: {str(e)}")
            traceback.print_exc()
            return jsonify({'success': False, 'query': query, 'patterns': [], 'categories': {}, 'error': str(e)}), 200
    
    except Exception as e:
        import traceback
        print(f"Error: {str(e)}")
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 200


@app.route('/api/graph-traversal', methods=['POST'])
def graph_traversal():
    """
    Execute graph traversal to find indirect connections.
    
    Body: {
        "start_entity": "Mumbai Port",
        "start_type": "Location",
        "target_type": "Organization" (optional),
        "max_depth": 3 (optional)
    }
    """
    try:
        data = request.get_json() or {}
        start_entity = data.get('start_entity', '')
        start_type = data.get('start_type', 'Location')
        target_type = data.get('target_type', 'Organization')
        max_depth = int(data.get('max_depth', 3))
        
        if not start_entity:
            return jsonify({'success': False, 'error': 'No start entity specified'}), 400
        
        try:
            from semantic_search.graph_traversal import GraphTraversal
            graph = GraphTraversal()
            
            result = graph.find_indirect_connections(
                start_entity, start_type, target_type, max_depth
            )
            
            # Format paths for frontend GraphPath type: { nodes, edges, depth }
            formatted_paths = []
            raw_paths = result.paths[:20] if hasattr(result, 'paths') else []
            for p in raw_paths:
                if isinstance(p, dict):
                    formatted_paths.append({
                        'nodes': p.get('nodes', []),
                        'edges': p.get('edges', []),
                        'depth': p.get('depth', len(p.get('nodes', [])) - 1)
                    })
                elif isinstance(p, (list, tuple)):
                    formatted_paths.append({
                        'nodes': list(p),
                        'edges': [],
                        'depth': len(p) - 1
                    })

            return jsonify({
                'success': True,
                'start_entity': start_entity,
                'start_type': start_type,
                'target_type': target_type,
                'depth': result.depth,
                'entities_found': result.entities,
                'relationships': result.relationships[:50],
                'paths': formatted_paths
            }), 200
            
        except Exception as e:
            print(f"Graph traversal error: {str(e)}")
            return jsonify({'success': False, 'error': str(e)}), 500
    
    except Exception as e:
        print(f"Error: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ingest', methods=['POST'])
def ingest_document():
    """
    Ingest a document through the semantic graph pipeline.
    
    Body: {
        "text": "Company X faced shipment delays...",
        "document_id": "doc_001" (optional),
        "category": "supply_chain" (optional),
        "source_type": "report" (optional)
    }
    """
    try:
        data = request.get_json() or {}
        text = data.get('text', data.get('content', ''))  # Support both 'text' and 'content' field names
        document_id = data.get('document_id', f"doc_{uuid.uuid4().hex[:8]}")
        category = data.get('category', 'general')
        source_type = data.get('source_type', 'document')
        
        if not text:
            return jsonify({'success': False, 'error': 'No text provided'}), 400
        
        try:
            from semantic_search.semantic_pipeline import SemanticGraphPipeline
            pipeline = SemanticGraphPipeline()
            
            result = pipeline.ingest_document(
                text=text,
                document_id=document_id,
                source_type=source_type,
                category=category
            )
            
            return jsonify({
                'success': True,
                'document_id': result.document_id,
                'summary': result.summary,
                'entities_count': len(result.entities),
                'relationships_count': len(result.relationships),
                # Both field name variants for frontend compatibility
                'entities_extracted': len(result.entities),
                'relationships_generated': len(result.relationships),
                'entities': result.entities,
                'relationships': result.relationships,
                'weaviate_stored': result.weaviate_id is not None,
                'stored_in_weaviate': result.weaviate_id is not None,
                'graph_stored': result.graph_stored,
                'stored_in_nebula': result.graph_stored
            }), 200
            
        except Exception as e:
            print(f"Ingest error: {str(e)}")
            return jsonify({'success': False, 'error': str(e)}), 500
    
    except Exception as e:
        print(f"Error: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/graph', methods=['GET'])
def get_graph_data():
    """Get entities and relationships from NebulaGraph"""
    try:
        # Use the standalone nebula handler to get graph data
        try:
            entities = nebula_handler.get_all_entities()
            relationships = nebula_handler.get_all_relationships()
            return jsonify({
                'success': True,
                'nodes': entities,
                'edges': relationships
            }), 200
        except Exception as e:
            print(f"NebulaGraph query error: {str(e)}")
            return jsonify({'success': True, 'nodes': [], 'edges': []}), 200
    
    except Exception as e:
        print(f"Graph data error: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/document/<document_id>/graph', methods=['GET'])
def get_document_graph(document_id):
    """Get graph data for a specific document"""
    try:
        import requests as http_requests

        # Search for the document in Weaviate
        graphql_query = {
            "query": f"""
            {{
                Get {{
                    SemanticDocument(
                        where: {{
                            path: ["document_id"],
                            operator: Equal,
                            valueText: "{document_id}"
                        }}
                    ) {{
                        document_id
                        content
                        category
                        _additional {{
                            id
                        }}
                    }}
                }}
            }}
            """
        }

        response = http_requests.post(
            "http://localhost:8080/v1/graphql",
            json=graphql_query,
            timeout=10
        )

        entities = []
        relationships = []
        doc_content = ""

        if response.status_code == 200:
            data = response.json()
            docs = data.get("data", {}).get("Get", {}).get("SemanticDocument", [])
            if docs:
                doc_content = docs[0].get("content", "")

        # Extract entities from the document content
        if doc_content:
            try:
                from semantic_search.semantic_engine import SemanticSearchEngine
                from semantic_search.relationship_mapper import EnhancedRelationshipMapper

                engine = SemanticSearchEngine()
                mapper = EnhancedRelationshipMapper()

                # Search for the document to get its entities
                results = engine.semantic_search(doc_content[:200], limit=1)
                if results:
                    for i, e in enumerate(results[0].entities):
                        entities.append({
                            'id': f'entity-{i}',
                            'name': e.get('value', e.get('name', '')),
                            'type': e.get('type', 'CUSTOM').upper(),
                            'confidence': e.get('confidence', 0.85)
                        })

                    doc_rels = mapper.map_relationships_for_entities(
                        results[0].entities, doc_content, ""
                    )
                    for j, rel in enumerate(doc_rels[:20]):
                        relationships.append({
                            'id': f'rel-{j}',
                            'source': rel.source,
                            'source_type': rel.source_type.upper(),
                            'target': rel.target,
                            'target_type': rel.target_type.upper(),
                            'type': rel.relationship_type,
                            'confidence': rel.confidence
                        })
            except Exception as e:
                print(f"Entity extraction for document graph failed: {e}")

        return jsonify({
            'document_id': document_id,
            'entities': entities,
            'relationships': relationships,
            'graph_data': {
                'nodes': [{'id': e['id'], 'name': e['name'], 'type': e['type']} for e in entities],
                'edges': [{'source': r['source'], 'target': r['target'], 'type': r['type']} for r in relationships]
            }
        }), 200

    except Exception as e:
        print(f"Document graph error: {str(e)}")
        return jsonify({
            'document_id': document_id,
            'entities': [],
            'relationships': [],
            'error': str(e)
        }), 500


@app.route('/api/graph/stats', methods=['GET'])
def get_graph_stats():
    """Get statistics from NebulaGraph"""
    try:
        try:
            stats = nebula_handler.get_graph_stats()
            return jsonify({
                'success': True,
                'stats': stats
            }), 200
        except Exception as e:
            print(f"NebulaGraph stats error: {str(e)}")
            return jsonify({'success': True, 'stats': {'entities': 0, 'relationships': 0}}), 200
    
    except Exception as e:
        print(f"Stats error: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/clear-all-data', methods=['POST'])
def clear_all_data():
    """
    Clear ALL stored documents and graph data from Weaviate and NebulaGraph.
    This is a full reset of all stored data.
    """
    try:
        results = {
            "weaviate": {"deleted": 0, "success": False},
            "nebula": {"deleted_vertices": 0, "deleted_edges": 0, "success": False},
        }

        # 1. Clear Weaviate
        try:
            from vector_database.weaviate_handler import WeaviateClient
            wv = WeaviateClient()
            if wv.client:
                wv_result = wv.delete_all_documents()
                results["weaviate"] = wv_result
                print(f"Weaviate cleared: {wv_result}")
        except Exception as e:
            results["weaviate"]["error"] = str(e)
            print(f"Weaviate clear error: {e}")

        # 2. Clear NebulaGraph
        try:
            from knowledge_graph.nebula_handler import NebulaGraphClient
            ng = NebulaGraphClient()
            ng_result = ng.clear_all_data()
            results["nebula"] = ng_result
            print(f"NebulaGraph cleared: {ng_result}")
        except Exception as e:
            results["nebula"]["error"] = str(e)
            print(f"NebulaGraph clear error: {e}")

        # 3. Clear uploaded files
        uploads_dir = os.path.join(os.path.dirname(__file__), 'uploads')
        files_deleted = 0
        if os.path.exists(uploads_dir):
            for f in os.listdir(uploads_dir):
                fpath = os.path.join(uploads_dir, f)
                if os.path.isfile(fpath):
                    os.remove(fpath)
                    files_deleted += 1

        total_deleted = results["weaviate"].get("deleted", 0) + results["nebula"].get("deleted_vertices", 0) + results["nebula"].get("deleted_edges", 0)

        return jsonify({
            "success": True,
            "message": f"Cleared all data: {total_deleted} items removed, {files_deleted} files deleted",
            "details": results,
            "files_deleted": files_deleted,
        }), 200

    except Exception as e:
        print(f"Clear all data error: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/clear-demo-data', methods=['POST'])
def clear_demo_data():
    """
    Clear demo/sample documents from Weaviate to get cleaner search results.
    This removes documents with IDs like 'doc_001', 'doc_002', etc.
    """
    try:
        import requests
        
        deleted_count = 0
        demo_ids = ['doc_001', 'doc_002', 'doc_003', 'doc_004', 'doc_005']
        
        # Delete from SemanticDocument class
        for doc_id in demo_ids:
            try:
                # Find documents with matching document_id
                graphql_query = {
                    "query": f"""
                    {{
                        Get {{
                            SemanticDocument(
                                where: {{
                                    path: ["document_id"],
                                    operator: Equal,
                                    valueText: "{doc_id}"
                                }}
                            ) {{
                                _additional {{
                                    id
                                }}
                            }}
                        }}
                    }}
                    """
                }
                
                response = requests.post(
                    "http://localhost:8080/v1/graphql",
                    json=graphql_query,
                    timeout=10
                )
                
                if response.status_code == 200:
                    data = response.json()
                    docs = data.get("data", {}).get("Get", {}).get("SemanticDocument", [])
                    for doc in docs:
                        uuid = doc.get("_additional", {}).get("id")
                        if uuid:
                            del_response = requests.delete(
                                f"http://localhost:8080/v1/objects/SemanticDocument/{uuid}",
                                timeout=10
                            )
                            if del_response.status_code in [200, 204]:
                                deleted_count += 1
                                print(f"Deleted demo document: {doc_id} (UUID: {uuid})")
            except Exception as e:
                print(f"Error deleting {doc_id}: {e}")
        
        # Also delete from ExtractedDocument class (legacy)
        for doc_id in demo_ids:
            try:
                graphql_query = {
                    "query": f"""
                    {{
                        Get {{
                            ExtractedDocument(
                                where: {{
                                    path: ["documentId"],
                                    operator: Equal,
                                    valueText: "{doc_id}"
                                }}
                            ) {{
                                _additional {{
                                    id
                                }}
                            }}
                        }}
                    }}
                    """
                }
                
                response = requests.post(
                    "http://localhost:8080/v1/graphql",
                    json=graphql_query,
                    timeout=10
                )
                
                if response.status_code == 200:
                    data = response.json()
                    docs = data.get("data", {}).get("Get", {}).get("ExtractedDocument", [])
                    for doc in docs:
                        uuid = doc.get("_additional", {}).get("id")
                        if uuid:
                            del_response = requests.delete(
                                f"http://localhost:8080/v1/objects/ExtractedDocument/{uuid}",
                                timeout=10
                            )
                            if del_response.status_code in [200, 204]:
                                deleted_count += 1
            except:
                pass
        
        return jsonify({
            'success': True,
            'message': f'Cleared {deleted_count} demo documents',
            'deleted_count': deleted_count
        }), 200
        
    except Exception as e:
        print(f"Error clearing demo data: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/delete-duplicates', methods=['POST'])
def delete_duplicates():
    """
    Find and delete duplicate documents from Weaviate.
    Duplicates are identified by having the same content (hash).
    Keeps the oldest document of each unique content.
    """
    try:
        import requests
        import hashlib
        
        deleted_count = 0
        docs_by_hash: dict = {}  # content_hash -> list of (uuid, created_time)
        
        # Get all documents from SemanticDocument class
        graphql_query = {
            "query": """
            {
                Get {
                    SemanticDocument(limit: 500) {
                        document_id
                        content
                        _additional {
                            id
                            creationTimeUnix
                        }
                    }
                }
            }
            """
        }
        
        response = requests.post(
            "http://localhost:8080/v1/graphql",
            json=graphql_query,
            timeout=30
        )
        
        if response.status_code == 200:
            data = response.json()
            docs = data.get("data", {}).get("Get", {}).get("SemanticDocument", [])
            
            # Group documents by content hash
            for doc in docs:
                content = doc.get("content", "") or ""
                content_hash = hashlib.md5(content.encode()).hexdigest()
                uuid = doc.get("_additional", {}).get("id", "")
                created = doc.get("_additional", {}).get("creationTimeUnix", 0)
                
                if content_hash not in docs_by_hash:
                    docs_by_hash[content_hash] = []
                docs_by_hash[content_hash].append({
                    "uuid": uuid,
                    "created": created,
                    "doc_id": doc.get("document_id", "")
                })
            
            # Delete duplicates (keep oldest)
            for content_hash, doc_list in docs_by_hash.items():
                if len(doc_list) > 1:
                    # Sort by creation time (oldest first)
                    sorted_docs = sorted(doc_list, key=lambda x: x["created"])
                    # Delete all but the first (oldest)
                    for dup in sorted_docs[1:]:
                        try:
                            del_response = requests.delete(
                                f"http://localhost:8080/v1/objects/SemanticDocument/{dup['uuid']}",
                                timeout=10
                            )
                            if del_response.status_code in [200, 204]:
                                deleted_count += 1
                                print(f"Deleted duplicate: {dup['doc_id']} (UUID: {dup['uuid']})")
                        except Exception as e:
                            print(f"Error deleting duplicate {dup['uuid']}: {e}")
        
        # Also check ExtractedDocument class
        graphql_query2 = {
            "query": """
            {
                Get {
                    ExtractedDocument(limit: 500) {
                        documentId
                        content
                        _additional {
                            id
                            creationTimeUnix
                        }
                    }
                }
            }
            """
        }
        
        response2 = requests.post(
            "http://localhost:8080/v1/graphql",
            json=graphql_query2,
            timeout=30
        )
        
        if response2.status_code == 200:
            data2 = response2.json()
            docs2 = data2.get("data", {}).get("Get", {}).get("ExtractedDocument", [])
            docs_by_hash2: dict = {}
            
            for doc in docs2:
                content = doc.get("content", "") or ""
                content_hash = hashlib.md5(content.encode()).hexdigest()
                uuid = doc.get("_additional", {}).get("id", "")
                created = doc.get("_additional", {}).get("creationTimeUnix", 0)
                
                if content_hash not in docs_by_hash2:
                    docs_by_hash2[content_hash] = []
                docs_by_hash2[content_hash].append({
                    "uuid": uuid,
                    "created": created,
                    "doc_id": doc.get("documentId", "")
                })
            
            for content_hash, doc_list in docs_by_hash2.items():
                if len(doc_list) > 1:
                    sorted_docs = sorted(doc_list, key=lambda x: x["created"])
                    for dup in sorted_docs[1:]:
                        try:
                            del_response = requests.delete(
                                f"http://localhost:8080/v1/objects/ExtractedDocument/{dup['uuid']}",
                                timeout=10
                            )
                            if del_response.status_code in [200, 204]:
                                deleted_count += 1
                                print(f"Deleted duplicate (legacy): {dup['doc_id']} (UUID: {dup['uuid']})")
                        except:
                            pass
        
        return jsonify({
            'success': True,
            'message': f'Removed {deleted_count} duplicate documents',
            'deleted_count': deleted_count
        }), 200
        
    except Exception as e:
        print(f"Error deleting duplicates: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/list-documents', methods=['GET'])
def list_documents():
    """List all documents stored in Weaviate for debugging"""
    try:
        import requests
        
        graphql_query = {
            "query": """
            {
                Get {
                    SemanticDocument(limit: 50) {
                        document_id
                        category
                        source_type
                        content
                        _additional {
                            id
                            creationTimeUnix
                        }
                    }
                }
            }
            """
        }
        
        response = requests.post(
            "http://localhost:8080/v1/graphql",
            json=graphql_query,
            timeout=15
        )
        
        if response.status_code == 200:
            data = response.json()
            docs = data.get("data", {}).get("Get", {}).get("SemanticDocument", [])
            
            # Format for display
            formatted = []
            for doc in docs:
                content = doc.get("content", "")[:100] + "..." if doc.get("content") else ""
                formatted.append({
                    "id": doc.get("_additional", {}).get("id", ""),
                    "document_id": doc.get("document_id", "unknown"),
                    "category": doc.get("category", ""),
                    "source_type": doc.get("source_type", ""),
                    "content_preview": content,
                    "created": doc.get("_additional", {}).get("creationTimeUnix", "")
                })
            
            return jsonify({
                'success': True,
                'documents': formatted,
                'total': len(formatted)
            }), 200
        
        return jsonify({'success': False, 'error': 'Failed to query Weaviate'}), 500
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


if __name__ == '__main__':
    print('=' * 60, flush=True)
    print('  Data Extraction Server', flush=True)
    print('  API:       http://localhost:5000/api/', flush=True)
    print('  Frontend:  http://localhost:5173', flush=True)
    print('=' * 60, flush=True)
    _log('Waiting for requests...')
    app.run(debug=False, port=5000, use_reloader=False)
